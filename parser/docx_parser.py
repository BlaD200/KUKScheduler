# Copyright (c) 2021 Vladyslav Synytsyn.

import re

from docx import Document as CreateDocument
from docx.document import Document
from docx.table import Table

from .parser_data import *


_teacher_regex = re.compile(
    r'(?P<teacher>'
    r'((([Дд]оц(ент)?\.? ?(з/н)?)|((ст\.)?[Вв]икл)|([Пп]роф)|[А-ЯІЇ][а-яії]{4,}) ?\.? ?'
    r'([ \w\-.а-яА-ЯіїІЇ])+?\.\10?.))')
_classroom_regex = re.compile(r'ауд.\s?(?P<classroom>\d+)')
_group_regex = re.compile(r'(?P<group>група\s*\d)(.*)?')
_date_regex = re.compile(r'(?P<date>\d{1,2}\.\d{1,2})\s?\(?(?P<classroom>\d{3})?')
_faculty_regex = re.compile(r'(?P<faculty>факультет\s*[а-яії ]*)', re.IGNORECASE)
_course_regex = re.compile(r'(?P<course>[\wі]+)\s?курс', re.IGNORECASE)
_year_regex = re.compile(
    r'((?P<semester>[\wі]+)\s?семестр)?\s*(?P<year_start>[\d]+)\s?[-–—]?\s?(?P<year_end>[\d]+)\s?н\.?р\.?',
    re.IGNORECASE
)


def _extract_edu_program_info(docx: Document):
    faculty: str | None = None
    course: str | None = None
    semester: str | None = None
    year_start: int | None = None
    year_end: int | None = None

    for i, p in enumerate(docx.paragraphs):
        print(i, p.text)
        if faculty_match := re.search(_faculty_regex, p.text):
            faculty = faculty_match.group("faculty")
        if course_match := re.search(_course_regex, p.text):
            course = course_match.group("course")
        if year_match := re.search(_year_regex, p.text):
            semester = year_match.group("semester")
            year_start = int(year_match.group("year_start")) if year_match.group("year_start") else None
            year_end = int(year_match.group("year_end")) if year_match.group("year_end") else None

        if faculty and course and year_start and year_end:
            break
    return faculty, course, semester, '%d-%d' % (year_start, year_end)


def _extract_subject_group(text: str, group: str = None) -> LessonGroup:
    teacher_match = re.search(_teacher_regex, text)
    teacher = teacher_match.group("teacher")

    classroom_match = re.search(_classroom_regex, text)
    if classroom_match:
        classroom = int(classroom_match.group("classroom"))
    else:
        classroom = None

    return LessonGroup(teacher, classroom, group)


def _extract_dates(text: str) -> list[str]:
    return [
        x.group('date')  # + (x.group('classroom') if x.group('classroom') else '')
        for x in re.finditer(_date_regex, text)
    ]


def _extract_tables_data_from_file(tables: list[Table]) -> list[list[list[str]]]:
    data_tables = []
    # table: Table
    for table in tables:
        df = [
            ['' for _ in range(len(table.columns))] for _ in range(len(table.rows))
        ]
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if cell.text.strip():
                    df[i][j] = cell.text.strip()  # '{:^15}'.format(cell.text.strip().split('\n')[0][:15])
                else:
                    df[i][j] = '{:^15}'.format('')
                # rels = cell.part.rels
                # for rel in cell.part.related_parts:
                #     if rels[rel].reltype == RELATIONSHIP_TYPE.HYPERLINK:
                #         # pass
                #         df[i][j] += rels[rel]._target
        data_tables.append(df)
    return data_tables


def parse_timetable():
    docx = CreateDocument('./test_data/1 курс ж, пр, мв КУК.docx')
    data_tables = _extract_tables_data_from_file(docx.tables)

    faculty, course, semester, studying_years = _extract_edu_program_info(docx)
    timetable_file = TimetableFile(faculty, course, studying_years, [])

    week_day_map = {
        0: 'monday',
        1: 'tuesday',
        2: 'wednesday',
        3: 'thursday',
        4: 'friday',
        5: 'saturday',
        6: 'sunday'
    }

    groups_map = {}
    for cell_idx, group_name in enumerate(data_tables[0][0]):
        if cell_idx < 2:
            continue
        group = Group(group_name, [])
        if group not in timetable_file.groups:
            timetable_file.groups.append(group)
            groups_map[cell_idx] = group

    # iterating through timetables for each day
    for i, table in enumerate(data_tables[:]):
        print('День ', i + 1)
        print('-' * (40 * 2 + 10 * 2 + 10))

        day = week_day_map[i]
        day_schedules_map = {}
        for cell_idx in groups_map:
            schedule = DaySchedule(day, [])
            groups_map[cell_idx].day_schedules.append(schedule)
            day_schedules_map[cell_idx] = schedule

        # row is a list of str in the format
        # first row: ['Пара', 'Тривалість', <group 1>, ..., <group N>]
        # other rows: [<lesson_slot>, <slot_duration>, <lesson 1>, ..., <lesson N>]
        for row in table[1:]:
            lesson_slot = int(row[0])
            lesson_slot_duration = re.sub(r'\s', '', row[1])

            for cell_idx, cell in enumerate(row[:]):
                if cell_idx < 2:
                    print('{:^11}'.format(''.join(cell.split('\n')[:3])), end=' | ')
                else:
                    print('{:^20}'.format(' '.join(cell.split('\n')[:3]))[:20], end=' | ')
                    if cell.strip():
                        lesson_groups = []
                        group_matchers = list(re.finditer(_group_regex, cell))
                        if group_matchers:
                            for group_match in group_matchers:
                                group = group_match.group('group')
                                group_srt = group_match.group()
                                lesson_group = _extract_subject_group(group_srt, group)
                                lesson_groups.append(lesson_group)
                        else:
                            lesson_group = _extract_subject_group(cell)
                            lesson_groups.append(lesson_group)

                        lines = cell.split('\n')
                        subject_name = lines[0]
                        lesson_type = lines[1].split(',')[0].split()[0]
                        dates = _extract_dates(cell)
                        lesson = Lesson(
                            subject_name,
                            lesson_type, lesson_slot, lesson_slot_duration,
                            dates, lesson_groups
                        )

                        if lesson not in day_schedules_map[cell_idx].lessons:
                            day_schedules_map[cell_idx].lessons.append(lesson)

            print('\n', '-' * (40 * 2 + 10 * 2 + 10))
        print()

        # Назва
        # Дати (Аудиотрія) дати
        # (Зум) .part.rels[].reltype == RT.HYPERLINK ? : rels[rel]._target
        # Викладач
        # (Зум)
        # Аудиторія (Optional)
        #
    return timetable_file


if __name__ == '__main__':
    timetable = parse_timetable()

# TODO use logging instead of print
